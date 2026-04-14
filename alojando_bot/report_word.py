"""
Generador de informe Word (.docx) usando python-docx.
Fallback: si Node.js + docx-js está disponible, usa el script JS.
"""
import os
import json
import subprocess
import logging
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from .models import ComparisonResult
from .config import OUTPUT_DIR

logger = logging.getLogger(__name__)

# Colores
PRIMARY = RGBColor(37, 99, 235)      # #2563EB
PRIMARY_DARK = RGBColor(30, 64, 175)  # #1E40AF
SUCCESS = RGBColor(22, 163, 74)       # #16A34A
WARNING = RGBColor(217, 119, 6)       # #D97706
DANGER = RGBColor(220, 38, 38)        # #DC2626
GRAY = RGBColor(148, 163, 184)        # #94A3B8
TEXT_COLOR = RGBColor(30, 41, 59)      # #1E293B


def generate_word_report(result: ComparisonResult, output_path: str = None) -> str:
    """
    Genera un informe Word (.docx).

    Args:
        result: ComparisonResult con el análisis
        output_path: Ruta para guardar el archivo .docx

    Returns:
        Ruta del archivo generado
    """
    if output_path is None:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(OUTPUT_DIR, f"informe_{timestamp}.docx")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    original = result.original
    currency = original.currency if original else "USD"

    doc = Document()

    # Configurar estilos
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT_COLOR

    # ============================================
    # PORTADA
    # ============================================
    for _ in range(6):
        doc.add_paragraph("")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("ALOJANDO BOT")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_DARK

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Informe Comparativo de Mercado")
    run.font.size = Pt(18)
    run.font.color.rgb = GRAY

    prop_p = doc.add_paragraph()
    prop_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prop_p.add_run(original.title or "Tu propiedad")
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(
        f"Generado el {datetime.now().strftime('%d/%m/%Y')} | "
        f"{len(result.comparables)} comparables analizados"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.add_page_break()

    # ============================================
    # RESUMEN EJECUTIVO
    # ============================================
    _add_heading(doc, "Resumen Ejecutivo")

    location = ", ".join(filter(None, [original.address, original.city, original.country]))

    _add_key_value(doc, "Propiedad: ", original.title or "No especificado")
    _add_key_value(doc, "Ubicacion: ", location or "No especificada")
    _add_key_value(doc, "Tu precio: ", f"{currency} {original.price_per_night:.0f}/noche")
    _add_key_value(doc, "Comparables encontrados: ", str(len(result.comparables)))
    _add_key_value(doc, "Precio mediana del mercado: ", f"{currency} {result.median_price:.0f}/noche")
    _add_key_value(doc, "Rating promedio zona: ", f"{result.avg_rating:.1f}/5.0")

    if result.suggested_price_low > 0 and result.suggested_price_high > 0:
        doc.add_paragraph("")
        highlight_p = doc.add_paragraph()
        run = highlight_p.add_run(
            f"Rango de precio sugerido: {currency} {result.suggested_price_low:.0f} - "
            f"{currency} {result.suggested_price_high:.0f}/noche"
        )
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_DARK
        # Add shading to paragraph
        shading = highlight_p.paragraph_format
        pPr = highlight_p._p.get_or_add_pPr()
        shd = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "DBEAFE"
        })
        pPr.append(shd)

    doc.add_page_break()

    # ============================================
    # ANALISIS DE PRECIOS
    # ============================================
    _add_heading(doc, "Analisis de Precios")

    # Tabla de estadísticas
    _add_subheading(doc, "Estadisticas de Precios")

    price_data = [
        ["Metrica", "Valor"],
        ["Tu precio", f"{currency} {original.price_per_night:.0f}"],
        ["Precio promedio mercado", f"{currency} {result.avg_price:.0f}"],
        ["Precio mediana mercado", f"{currency} {result.median_price:.0f}"],
        ["Precio minimo", f"{currency} {result.min_price:.0f}"],
        ["Precio maximo", f"{currency} {result.max_price:.0f}"],
        ["Tu percentil", f"{result.price_percentile:.0f}%"],
    ]
    _add_table(doc, price_data)

    _add_subheading(doc, "Sugerencias de Precio")
    for s in result.pricing_suggestions:
        _add_bullet(doc, s)

    doc.add_page_break()

    # ============================================
    # TABLA DE COMPARABLES
    # ============================================
    _add_heading(doc, "Detalle de Comparables")

    comp_data = [["Portal", "Titulo", "Precio", "Rating", "Dorm."]]

    # Original primero
    comp_data.append([
        (original.source or "TU ANUNCIO").upper(),
        (original.title or "Tu propiedad")[:40],
        f"{currency} {original.price_per_night:.0f}",
        f"{original.rating:.1f}" if original.rating > 0 else "-",
        str(original.bedrooms) if original.bedrooms > 0 else "-"
    ])

    for c in result.comparables[:20]:
        comp_data.append([
            c.source.upper(),
            (c.title or "")[:40],
            f"{currency} {c.price_per_night:.0f}" if c.price_per_night > 0 else "-",
            f"{c.rating:.1f}" if c.rating > 0 else "-",
            str(c.bedrooms) if c.bedrooms > 0 else "-"
        ])

    _add_table(doc, comp_data, col_widths=[Cm(2.5), Cm(6), Cm(2.5), Cm(2), Cm(2)])

    doc.add_page_break()

    # ============================================
    # AMENIDADES
    # ============================================
    _add_heading(doc, "Analisis de Amenidades")

    if result.common_amenities:
        _add_subheading(doc, "Amenidades mas comunes en la zona")
        for amenity, count in result.common_amenities:
            _add_bullet(doc, f"{amenity} ({count} propiedades)")

    if result.missing_amenities:
        _add_subheading(doc, "Amenidades que te faltan")
        for a in result.missing_amenities:
            _add_bullet(doc, a)

    if result.unique_amenities:
        _add_subheading(doc, "Tus amenidades unicas (ventaja competitiva)")
        for a in result.unique_amenities:
            _add_bullet(doc, a)

    if result.amenity_suggestions:
        _add_subheading(doc, "Sugerencias")
        for s in result.amenity_suggestions:
            _add_bullet(doc, s)

    doc.add_page_break()

    # ============================================
    # CALIFICACIONES
    # ============================================
    _add_heading(doc, "Calificaciones y Resenas")
    if result.rating_comparison:
        p = doc.add_paragraph(result.rating_comparison)
        p.style.font.size = Pt(11)

    # ============================================
    # SUGERENCIAS DE TITULO
    # ============================================
    doc.add_paragraph("")
    _add_heading(doc, "Mejoras en Titulo")
    for s in result.title_suggestions:
        _add_bullet(doc, s)

    # ============================================
    # SUGERENCIAS DE DESCRIPCION
    # ============================================
    doc.add_page_break()
    _add_heading(doc, "Mejoras en Descripcion")
    for s in result.description_suggestions:
        _add_bullet(doc, s)

    # ============================================
    # SUGERENCIAS DE FOTOS
    # ============================================
    doc.add_paragraph("")
    _add_heading(doc, "Mejoras en Fotos")
    for s in result.photo_suggestions:
        _add_bullet(doc, s)

    # ============================================
    # RECOMENDACIONES GENERALES
    # ============================================
    doc.add_page_break()
    _add_heading(doc, "Recomendaciones Generales")
    for s in result.general_suggestions:
        _add_bullet(doc, s)

    # ============================================
    # PIE DE PAGINA
    # ============================================
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        "Este informe es orientativo. Los datos provienen de fuentes publicas y pueden variar."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    # Guardar
    doc.save(output_path)
    logger.info(f"Informe Word generado: {output_path}")
    return output_path


def _add_heading(doc, text):
    """Agrega un heading estilizado."""
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = PRIMARY_DARK
        run.font.name = "Arial"
    # Línea divisoria
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "12",
        qn("w:space"): "1",
        qn("w:color"): "2563EB"
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_subheading(doc, text):
    """Agrega un subheading."""
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = PRIMARY
        run.font.name = "Arial"


def _add_key_value(doc, label, value):
    """Agrega un par clave: valor."""
    p = doc.add_paragraph()
    run_label = p.add_run(label)
    run_label.font.bold = True
    run_label.font.size = Pt(11)
    run_label.font.name = "Arial"
    run_value = p.add_run(value)
    run_value.font.size = Pt(11)
    run_value.font.name = "Arial"


def _add_bullet(doc, text):
    """Agrega un bullet point."""
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Arial"


def _add_table(doc, data, col_widths=None):
    """Agrega una tabla estilizada."""
    if not data:
        return

    rows = len(data)
    cols = len(data[0])

    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text

            # Estilo del texto
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    elif i == 1 and rows > 2:
                        run.font.bold = True

            # Color de fondo
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "2563EB" if i == 0 else ("DBEAFE" if i == 1 and rows > 2 else ("F8FAFC" if i % 2 == 0 else "FFFFFF"))
            })
            shading.append(shd)

    # Aplicar anchos de columna si se especificaron
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    return table
